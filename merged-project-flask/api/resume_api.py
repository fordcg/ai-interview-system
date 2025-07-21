#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
简历生成API
提供AI简历生成、解析和实体识别功能
"""

import json
import base64
from flask import Blueprint, request, jsonify, current_app
import time
import requests
from datetime import datetime
from wsgiref.handlers import format_date_time
from time import mktime
import hashlib
import hmac
from urllib.parse import urlencode
import traceback
import os
import sys
from typing import List, Dict, Any, Optional
from werkzeug.utils import secure_filename

# 导入公共函数
from api.common import allowed_file

# 导入简历实体识别模块
from utils.resume_ner import get_resume_ner

# 创建蓝图
resumeBp = Blueprint('resume', __name__)

class AssembleHeaderException(Exception):
    def __init__(self, msg):
        self.message = msg

class Url:
    def __init__(self, host, path, schema):
        self.host = host
        self.path = path
        self.schema = schema

# calculate sha256 and encode to base64
def sha256base64(data):
    sha256 = hashlib.sha256()
    sha256.update(data)
    digest = base64.b64encode(sha256.digest()).decode(encoding='utf-8')
    return digest

def parse_url(requset_url):
    stidx = requset_url.index("://")
    host = requset_url[stidx + 3:]
    schema = requset_url[:stidx + 3]
    edidx = host.index("/")
    if edidx <= 0:
        raise AssembleHeaderException("invalid request url:" + requset_url)
    path = host[edidx:]
    host = host[:edidx]
    u = Url(host, path, schema)
    return u

# 生成鉴权url
def assemble_ws_auth_url(requset_url, method="GET", api_key="", api_secret=""):
    u = parse_url(requset_url)
    host = u.host
    path = u.path
    now = datetime.now()
    date = format_date_time(mktime(now.timetuple()))
    signature_origin = "host: {}\ndate: {}\n{} {} HTTP/1.1".format(host, date, method, path)
    signature_sha = hmac.new(api_secret.encode('utf-8'), signature_origin.encode('utf-8'),
                             digestmod=hashlib.sha256).digest()
    signature_sha = base64.b64encode(signature_sha).decode(encoding='utf-8')
    authorization_origin = "api_key=\"%s\", algorithm=\"%s\", headers=\"%s\", signature=\"%s\"" % (
        api_key, "hmac-sha256", "host date request-line", signature_sha)
    authorization = base64.b64encode(authorization_origin.encode('utf-8')).decode(encoding='utf-8')
    values = {
        "host": host,
        "date": date,
        "authorization": authorization
    }
    return requset_url + "?" + urlencode(values)

# 生成请求body体
def getBody(appid, text):
    body = {
        "header": {
            "app_id": appid,
            "status": 3,
        },
        "parameter": {
            "ai_resume": {
                "resData": {
                    "encoding": "utf8",
                    "compress": "raw",
                    "format": "json"
                }
            }
        },
        "payload": {
            "reqData": {
                "encoding": "utf8",
                "compress": "raw",
                "format": "plain",
                "status": 3,
                "text": base64.b64encode(text.encode("utf-8")).decode('utf-8')
            }
        }
    }
    return body

# 调用讯飞API生成简历
def generate_resume(text, appid, apikey, apisecret):
    host = 'https://cn-huadong-1.xf-yun.com/v1/private/s73f4add9'
    url = assemble_ws_auth_url(host, method='POST', api_key=apikey, api_secret=apisecret)
    content = getBody(appid, text)
    response = requests.post(url, json=content, headers={'content-type': "application/json"}).text
    return response

# 从URL下载文档并提取文本
def extract_text_from_url(document_url):
    """
    从URL下载文档并提取文本内容

    Args:
        document_url: 文档URL

    Returns:
        提取的文本内容，失败时返回None
    """
    try:
        print(f"开始从URL下载文档: {document_url}")

        # 下载文档
        response = requests.get(document_url, timeout=30)
        response.raise_for_status()

        # 创建临时文件
        import tempfile
        import os

        # 根据URL判断文件类型
        file_extension = '.docx'
        if document_url.lower().endswith('.pdf'):
            file_extension = '.pdf'
        elif document_url.lower().endswith('.doc'):
            file_extension = '.doc'

        # 保存到临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            temp_file.write(response.content)
            temp_file_path = temp_file.name

        print(f"文档已下载到临时文件: {temp_file_path}")

        # 提取文本
        try:
            # 导入文本提取模块
            sys.path.append(os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'scripts/workflow'))
            import document_text_extractor

            # 提取文本
            text_content = document_text_extractor.extract_text_from_document(temp_file_path)
            print(f"成功提取文本，长度: {len(text_content) if text_content else 0}")

            return text_content

        finally:
            # 清理临时文件
            try:
                os.unlink(temp_file_path)
                print(f"临时文件已删除: {temp_file_path}")
            except Exception as cleanup_error:
                print(f"删除临时文件失败: {cleanup_error}")

    except Exception as e:
        print(f"从URL提取文本失败: {str(e)}")
        return None

# 降级技能提取函数
def extract_skills_fallback(text):
    """
    当NER失败时使用的降级技能提取方法
    基于关键词匹配提取技能

    Args:
        text: 简历文本

    Returns:
        提取的技能列表
    """
    skills = set()
    text_lower = text.lower()

    # 常见技能关键词
    skill_keywords = [
        'python', 'java', 'javascript', 'php', 'c++', 'c#', 'go', 'rust',
        'html', 'css', 'react', 'vue', 'angular', 'node.js', 'express',
        'django', 'flask', 'spring', 'mysql', 'postgresql', 'mongodb',
        'redis', 'docker', 'kubernetes', 'aws', 'azure', 'git', 'linux',
        '机器学习', '深度学习', '数据分析', '人工智能', '大数据',
        '前端开发', '后端开发', '全栈开发', '移动开发', '测试',
        '项目管理', '团队协作', '沟通能力', '问题解决'
    ]

    for keyword in skill_keywords:
        if keyword in text_lower:
            skills.add(keyword)

    return list(skills)

# 使用RANER模型分析简历文本
def analyze_resume_text(text):
    """
    使用RANER模型分析简历文本，提取命名实体
    
    Args:
        text: 简历文本
    
    Returns:
        提取的结构化信息
    """
    try:
        # 检查是否启用了简历实体识别功能
        if not current_app.config.get('RESUME_NER_ENABLED', False):
            print("简历实体识别功能未启用")
            return None
        
        # 获取模型路径
        model_path = current_app.config.get('RESUME_NER_MODEL_PATH', None)
        
        # 获取NER实例
        resume_ner = get_resume_ner(model_path)
        
        # 提取结构化信息
        structured_info = resume_ner.extract_structured_info(text)
        
        print(f"简历分析结果: {structured_info}")
        return structured_info
    
    except Exception as e:
        print(f"简历分析失败: {str(e)}")
        traceback.print_exc()
        return None

# 从Word文档中提取文本
def extract_text_from_docx(file_path):
    """
    从Word文档中提取文本
    
    Args:
        file_path: Word文档的本地路径
        
    Returns:
        提取的文本内容，包括段落、表格和元数据
    """
    try:
        print(f"正在从Word文档提取文本: {file_path}")
        
        # 打开Word文档
        from docx import Document
        doc = Document(file_path)
        
        # 提取文档属性
        doc_properties = {
            'title': doc.core_properties.title or '',
            'author': doc.core_properties.author or '',
            'created': str(doc.core_properties.created) if doc.core_properties.created else '',
            'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
            'last_modified_by': doc.core_properties.last_modified_by or ''
        }
        
        # 提取段落文本
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # 提取表格文本
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # 获取单元格中的所有段落文本
                    cell_text = '\n'.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    table_data.append(' | '.join(row_data))
            if table_data:
                tables.append('\n'.join(table_data))
        
        # 合并所有文本
        all_text = '\n\n'.join(paragraphs + tables)
        
        return all_text, doc_properties
    
    except Exception as e:
        print(f"提取文本失败: {str(e)}")
        traceback.print_exc()
        return None, {}

# 清理文本
def clean_text(text):
    """
    清理文本，去除多余的空白字符和特殊字符
    
    Args:
        text: 待清理的文本
    
    Returns:
        清理后的文本
    """
    if not text:
        return ""
    
    # 替换多个连续的空白字符为单个空格
    import re
    text = re.sub(r'\s+', ' ', text)
    
    # 去除特殊字符
    text = re.sub(r'[^\w\s\u4e00-\u9fff,.，。:：;；!！?？、""''（）()【】[\]{}]', '', text)
    
    # 去除首尾空白
    text = text.strip()
    
    return text

# 解析简历API返回的数据，提取所有简历模板链接
def parse_resume_api_response(response_data: str) -> List[Dict[str, str]]:
    """
    解析简历API返回的数据，提取所有简历模板链接
    
    Args:
        response_data: API返回的JSON字符串
    
    Returns:
        简历模板列表，每个模板包含图片URL和Word文档URL
    """
    try:
        # 解析JSON数据
        result = json.loads(response_data)
        
        # 检查返回结果
        code = result['header']['code']
        if code != 0:
            error_msg = result['header'].get('message', '未知错误')
            print(f"API错误: {error_msg}")
            return []
        
        # 获取生成的简历内容
        text_base64 = result['payload']['resData']['text']
        print(f"Base64编码: {text_base64[:50]}...") # 只打印前50个字符
        
        # 解码
        resume_bytes = base64.b64decode(text_base64)
        resume_text = resume_bytes.decode('utf-8')
        print(f"解码后: {resume_text[:100]}...") # 只打印前100个字符
        
        # 解析JSON
        resume_json = json.loads(resume_text)
        print(f"解析JSON成功")
        
        # 检查是否包含links字段
        if 'links' in resume_json and isinstance(resume_json['links'], list) and len(resume_json['links']) > 0:
            return resume_json['links']
        else:
            print("未找到简历模板链接")
            return []
    
    except Exception as e:
        print(f"解析API返回数据失败: {str(e)}")
        return []

# 基于简历信息获取推荐职位
def get_recommended_jobs(resume_info: Dict[str, Any], api_url: str = None) -> List[Dict[str, Any]]:
    """
    基于简历信息获取推荐职位
    
    Args:
        resume_info: 简历信息，包含教育背景、技能等
        api_url: 职位推荐API的URL，如果为None则使用默认API
    
    Returns:
        推荐职位列表
    """
    try:
        # 如果没有提供API URL，使用默认的推荐API
        if api_url is None:
            api_url = "http://localhost:5000/job/getRecomendation"
        
        # 提取简历中的关键信息
        education = resume_info.get('education', '')
        skills = resume_info.get('skills', [])
        experience = resume_info.get('experience', '')
        
        # 构建请求参数
        params = {
            'education': education,
            'skills': ','.join(skills) if isinstance(skills, list) else skills,
            'experience': experience
        }
        
        # 发送请求获取推荐职位
        print(f"正在获取推荐职位...")
        response = requests.get(api_url, params=params)
        
        if response.status_code != 200:
            print(f"API请求失败，状态码: {response.status_code}")
            return []
        
        # 解析返回结果
        result = response.json()
        
        if result.get('code') != 0:
            print(f"获取推荐职位失败: {result.get('msg', '未知错误')}")
            return []
        
        # 返回推荐职位列表
        jobs = result.get('data', [])
        return jobs
    
    except Exception as e:
        print(f"获取推荐职位时发生错误: {str(e)}")
        return []

# 临时文件管理器
class TempFile:
    def __init__(self, file_path):
        self.file_path = file_path
        self.is_temp = False
    
    def __enter__(self):
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.is_temp and self.file_path and os.path.exists(self.file_path):
            try:
                os.remove(self.file_path)
                print(f"临时文件已删除: {self.file_path}")
            except Exception as e:
                print(f"删除临时文件失败: {str(e)}")

@resumeBp.route('/resume/generate', methods=['POST'])
def generate_ai_resume():
    """
    生成AI简历
    
    请求参数:
        name: 姓名
        education: 教育经历
        skills: 技能列表
        experience: 工作经历
    
    返回:
        生成的简历内容
    """
    try:
        data = request.get_json()
        
        # 获取请求参数
        name = data.get('name', '')
        education = data.get('education', '')
        skills = data.get('skills', [])
        experience = data.get('experience', '')
        
        # 构建简历文本
        if isinstance(skills, list):
            skills_text = '、'.join(skills)
        else:
            skills_text = skills
        
        resume_text = f"姓名：{name}\n教育经历：{education}\n技能：{skills_text}\n工作经历：{experience}"
        
        # 获取API配置
        appid = current_app.config.get('XF_APPID')
        apikey = current_app.config.get('XF_API_KEY')
        apisecret = current_app.config.get('XF_APISECRET')

        # 调试信息
        print(f"讯飞API配置检查:")
        print(f"APPID: {appid}")
        print(f"APIKEY: {apikey[:10] + '...' if apikey else None}")
        print(f"APISECRET: {apisecret[:10] + '...' if apisecret else None}")

        if not all([appid, apikey, apisecret]):
            missing_configs = []
            if not appid:
                missing_configs.append('XF_APPID')
            if not apikey:
                missing_configs.append('XF_API_KEY')
            if not apisecret:
                missing_configs.append('XF_APISECRET')

            return jsonify({
                'code': 500,
                'message': f'API配置不完整，缺少: {", ".join(missing_configs)}',
                'data': None
            })
        
        # 调用讯飞API生成简历
        print(f"正在生成简历，输入文本: {resume_text}")
        response = generate_resume(resume_text, appid, apikey, apisecret)
        
        # 解析API返回结果
        result = json.loads(response)
        code = result['header']['code']
        
        if code == 0:
            # 获取生成的简历内容
            text_base64 = result['payload']['resData']['text']
            
            try:
                # 解码
                resume_bytes = base64.b64decode(text_base64)
                resume_text = resume_bytes.decode('utf-8')
                
                # 解析JSON
                resume_json = json.loads(resume_text)
                
                # 提取实体 - 从生成的简历文档中提取
                entities = None
                if current_app.config.get('RESUME_NER_ENABLED', False):
                    try:
                        # 检查是否有简历链接可以下载和分析
                        document_text = None
                        if 'links' in resume_json and isinstance(resume_json['links'], list) and len(resume_json['links']) > 0:
                            # 尝试从第一个链接下载并提取文本
                            first_link = resume_json['links'][0]
                            if 'word_url' in first_link:
                                document_text = extract_text_from_url(first_link['word_url'])

                        if document_text and len(document_text.strip()) > 50:
                            # 对文档文本进行NER分析
                            print(f"对下载的文档文本进行NER分析，文本长度: {len(document_text)}")
                            entities = analyze_resume_text(document_text)
                        else:
                            print("未能获取有效的文档文本，跳过NER分析")
                            entities = None

                        # 确保entities可以JSON序列化
                        if entities:
                            entities = json.loads(json.dumps(entities, default=str))
                    except Exception as entity_error:
                        print(f"实体识别失败: {str(entity_error)}")
                        entities = None

                # 返回结果
                return jsonify({
                    'code': 0,
                    'message': '简历生成成功',
                    'data': {
                        'resume': resume_json,
                        'entities': entities  # 添加实体识别结果
                    }
                })
            except Exception as e:
                print(f"解码错误: {str(e)}")
                # 返回原始响应，不进行解码
                # 确保result可以JSON序列化
                try:
                    serializable_result = json.loads(json.dumps(result, default=str))
                except:
                    serializable_result = str(result)

                return jsonify({
                    'code': 0,
                    'message': '简历生成成功，但解码失败',
                    'data': {
                        'raw_response': serializable_result,
                        'error': str(e),
                        'input': {
                            'name': name,
                            'education': education,
                            'skills': skills,
                            'experience': experience
                        }
                    }
                })
        else:
            # 返回错误信息
            error_msg = result['header'].get('message', '未知错误')
            return jsonify({
                'code': code,
                'message': f'简历生成失败: {error_msg}',
                'data': None
            })
    
    except Exception as e:
        import traceback
        print(f"服务器错误: {str(e)}")
        print(traceback.format_exc())
        return jsonify({
            'code': 500,
            'message': f"服务器错误: {str(e)}",
            'data': None
        })

@resumeBp.route('/resume/analyze', methods=['POST'])
def analyze_resume():
    """
    分析简历文本，提取命名实体
    
    请求参数:
        text: 简历文本
    
    返回:
        提取的结构化信息
    """
    try:
        data = request.get_json()
        
        # 获取请求参数
        text = data.get('text', '')
        
        if not text:
            return jsonify({
                'code': 400,
                'message': '请求参数错误: 缺少text',
                'data': None
            })
        
        # 分析简历文本
        entities = analyze_resume_text(text)
        
        if entities is None:
            return jsonify({
                'code': 500,
                'message': '简历分析失败',
                'data': None
            })
        
        return jsonify({
            'code': 0,
            'message': '简历分析成功',
            'data': {
                'entities': entities
            }
        })
    
    except Exception as e:
        import traceback
        print(f"服务器错误: {str(e)}")
        print(traceback.format_exc())
        return jsonify({
            'code': 500,
            'message': f"服务器错误: {str(e)}",
            'data': None
        }) 

@resumeBp.route('/resume/process', methods=['POST'])
def process_resume_doc():
    """
    处理简历文档并提取实体
    
    请求参数:
        resumeUrl: 简历文档URL
    
    返回:
        提取的实体信息
    """
    temp_file_path = None
    
    try:
        data = request.get_json()
        
        # 获取请求参数
        resume_url = data.get('resumeUrl', '')
        
        # 使用临时文件上下文管理器
        with TempFile(None) as temp_file:
            if not resume_url:
                # 如果没有提供URL，使用默认文档路径
                doc_path = 'D:/zp/test_docs/f4a63f81-738b-4c54-ad18-88f1990fe52d.docx'
                print(f"未提供简历URL，使用默认文档: {doc_path}")
                
                # 检查文件是否存在
                if not os.path.exists(doc_path):
                    return jsonify({
                        'code': 404,
                        'message': f'文件不存在: {doc_path}',
                        'data': None
                    })
                
                temp_file.file_path = doc_path
            else:
                # 下载指定URL的文档
                print(f"准备下载简历文档: {resume_url}")
                
                # 创建临时文件
                import tempfile
                
                temp_dir = tempfile.gettempdir()
                file_name = f"resume_{int(time.time())}.docx"
                temp_file.file_path = os.path.join(temp_dir, file_name)
                temp_file.is_temp = True
                
                # 下载文件
                try:
                    # 设置超时和重试
                    from requests.adapters import HTTPAdapter
                    from requests.packages.urllib3.util.retry import Retry
                    
                    session = requests.Session()
                    retry = Retry(total=3, backoff_factor=0.5)
                    adapter = HTTPAdapter(max_retries=retry)
                    session.mount('http://', adapter)
                    session.mount('https://', adapter)
                    
                    response = session.get(resume_url, stream=True, timeout=(5, 30))  # 连接超时5秒，读取超时30秒
                    response.raise_for_status()  # 如果请求失败，抛出异常
                    
                    with open(temp_file.file_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            if chunk:
                                f.write(chunk)
                    
                    print(f"文档下载成功，保存至: {temp_file.file_path}")
                except requests.exceptions.RequestException as e:
                    print(f"下载文档失败: {str(e)}")
                    return jsonify({
                        'code': 500,
                        'message': f"下载文档失败: {str(e)}",
                        'data': None
                    })
            
            # 从Word文档提取文本
            text, doc_properties = extract_text_from_docx(temp_file.file_path)
            
            if not text:
                return jsonify({
                    'code': 500,
                    'message': "提取文本失败或文档为空",
                    'data': None
                })
            
            # 文本预处理：分段处理长文本
            max_text_length = 5000  # 最大文本长度
            if len(text) > max_text_length:
                print(f"文本过长({len(text)}字符)，将进行分段处理")
                # 简单分段策略：按段落分割
                segments = text.split('\n\n')
                processed_segments = []
                current_segment = ""
                
                for segment in segments:
                    if len(current_segment) + len(segment) < max_text_length:
                        current_segment += segment + '\n\n'
                    else:
                        processed_segments.append(current_segment)
                        current_segment = segment + '\n\n'
                
                if current_segment:
                    processed_segments.append(current_segment)
                
                # 处理每个分段并合并结果
                all_entities = []
                for i, segment in enumerate(processed_segments):
                    print(f"处理分段 {i+1}/{len(processed_segments)}")
                    segment_entities = analyze_resume_text(segment)
                    if segment_entities:
                        if not all_entities:
                            all_entities = segment_entities
                        else:
                            # 合并实体列表
                            for key, values in segment_entities.items():
                                if key == 'raw_entities':
                                    all_entities[key].extend(values)
                                elif isinstance(values, list):
                                    all_entities[key].extend(values)
                
                entities = all_entities
            else:
                # 使用RANER模型提取实体
                entities = analyze_resume_text(text)
            
            if entities is None:
                return jsonify({
                    'code': 500,
                    'message': '实体提取失败',
                    'data': None
                })
            
            # 返回结果
            return jsonify({
                'code': 0,
                'message': '简历处理成功',
                'data': {
                    'text': text[:1000] + ('...' if len(text) > 1000 else ''),  # 限制返回文本长度
                    'text_length': len(text),
                    'entities': entities,
                    'doc_properties': doc_properties
                }
            })
    
    except Exception as e:
        import traceback
        print(f"服务器错误: {str(e)}")
        print(traceback.format_exc())
        
        # 确保清理临时文件
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
                print(f"临时文件已删除: {temp_file_path}")
            except:
                pass
        
        return jsonify({
            'code': 500,
            'message': f"服务器错误: {str(e)}",
            'data': None
        })

@resumeBp.route('/resume/parse_templates', methods=['POST'])
def parse_resume_templates():
    """
    解析简历API返回的数据，提取所有简历模板链接
    
    请求参数:
        response_data: API返回的JSON字符串
    
    返回:
        简历模板列表，每个模板包含图片URL和Word文档URL
    """
    try:
        data = request.get_json()
        
        # 获取请求参数
        response_data = data.get('response_data', '')
        
        if not response_data:
            return jsonify({
                'code': 400,
                'message': '请求参数错误: 缺少response_data',
                'data': None
            })
        
        # 解析简历API返回的数据
        templates = parse_resume_api_response(response_data)
        
        if not templates:
            return jsonify({
                'code': 404,
                'message': '未找到简历模板',
                'data': None
            })
        
        return jsonify({
            'code': 0,
            'message': '解析成功',
            'data': {
                'templates': templates,
                'count': len(templates)
            }
        })
    
    except Exception as e:
        import traceback
        print(f"服务器错误: {str(e)}")
        print(traceback.format_exc())
        return jsonify({
            'code': 500,
            'message': f"服务器错误: {str(e)}",
            'data': None
        })

@resumeBp.route('/resume/recommend_jobs', methods=['POST'])
def recommend_jobs():
    """
    基于简历信息获取推荐职位
    
    请求参数:
        name: 姓名
        education: 教育经历
        skills: 技能列表
        experience: 工作经历
    
    返回:
        推荐职位列表
    """
    try:
        data = request.get_json()
        
        # 获取请求参数
        name = data.get('name', '')
        education = data.get('education', '')
        skills = data.get('skills', [])
        experience = data.get('experience', '')
        
        # 构建简历信息
        resume_info = {
            'name': name,
            'education': education,
            'skills': skills,
            'experience': experience
        }
        
        # 获取推荐职位API URL
        api_url = current_app.config.get('JOB_RECOMMENDATION_API_URL')
        
        # 获取推荐职位
        jobs = get_recommended_jobs(resume_info, api_url)
        
        if not jobs:
            return jsonify({
                'code': 404,
                'message': '未找到推荐职位',
                'data': None
            })
        
        return jsonify({
            'code': 0,
            'message': '获取推荐职位成功',
            'data': {
                'jobs': jobs,
                'count': len(jobs)
            }
        })
    
    except Exception as e:
        import traceback
        print(f"服务器错误: {str(e)}")
        print(traceback.format_exc())
        return jsonify({
            'code': 500,
            'message': f"服务器错误: {str(e)}",
            'data': None
        }) 

@resumeBp.route('/extract_text', methods=['POST'])
def extract_text():
    """从上传的文件中提取文本"""
    try:
        # 检查是否有文件
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'message': '没有文件部分',
                'data': None
            }), 400
        
        file = request.files['file']
        
        # 检查文件名是否为空
        if file.filename == '':
            return jsonify({
                'success': False,
                'message': '没有选择文件',
                'data': None
            }), 400
        
        # 检查文件类型是否允许
        if not allowed_file(file.filename):
            return jsonify({
                'success': False,
                'message': '文件类型不允许，请上传PDF或Word文档',
                'data': None
            }), 400
        
        # 保存文件
        filename = secure_filename(file.filename)
        # 添加时间戳前缀，确保文件名唯一
        unique_filename = f"{int(time.time())}_{filename}"
        file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)
        
        print(f"文件已保存为: {file_path}")
        
        # 提取文本
        try:
            # 导入文本提取模块
            sys.path.append(os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'scripts/workflow'))
            import document_text_extractor
            
            # 提取文本
            text_content = document_text_extractor.extract_text_from_document(file_path)
            
            if not text_content:
                return jsonify({
                    'success': False,
                    'message': '无法从文件中提取文本',
                    'data': None
                }), 500
            
            print(f"成功提取文本，长度: {len(text_content)} 字符")
            
            # 提取技能信息
            skills = []
            ner_success = False
            ner_error_message = None

            try:
                # 导入简历实体识别模块
                from utils.resume_ner import get_resume_ner

                # 获取NER实例
                resume_ner = get_resume_ner()

                print(f"开始NER处理，文本长度: {len(text_content)} 字符")

                # 提取结构化信息
                structured_info = resume_ner.extract_structured_info(text_content)

                # 获取技能信息
                if structured_info and 'skills' in structured_info:
                    skills = structured_info['skills']
                    print(f"成功提取技能: {skills}")
                    ner_success = True
                else:
                    print("NER处理完成，但未提取到技能信息")
                    ner_success = True  # NER成功，只是没有技能

            except Exception as e:
                ner_error_message = str(e)
                print(f"NER处理失败: {ner_error_message}")
                traceback.print_exc()

                # 降级处理：使用简单的关键词提取
                try:
                    print("尝试使用降级技能提取方法...")
                    skills = extract_skills_fallback(text_content)
                    if skills:
                        print(f"降级方法提取到技能: {skills}")
                except Exception as fallback_error:
                    print(f"降级技能提取也失败: {str(fallback_error)}")
                    skills = []
            
            # 构建响应消息
            if ner_success:
                message = '文本提取和技能识别成功'
            elif ner_error_message:
                message = f'文本提取成功，但NER处理失败: {ner_error_message}。已使用降级方法提取技能。'
            else:
                message = '文本提取成功，技能识别部分失败'

            return jsonify({
                'success': True,
                'message': message,
                'data': {
                    'filename': filename,
                    'file_path': file_path,
                    'text': text_content,
                    'text_length': len(text_content),
                    'skills': skills,  # 添加提取的技能信息
                    'ner_success': ner_success,  # NER处理是否成功
                    'ner_error': ner_error_message  # NER错误信息（如果有）
                }
            })
        
        except ImportError as e:
            print(f"导入文本提取模块失败: {str(e)}")
            return jsonify({
                'success': False,
                'message': f"导入文本提取模块失败: {str(e)}",
                'data': None
            }), 500
        
        except Exception as e:
            print(f"文本提取失败: {str(e)}")
            return jsonify({
                'success': False,
                'message': f"文本提取失败: {str(e)}",
                'data': None
            }), 500
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f"服务器错误: {str(e)}",
            'data': None
        }), 500

@resumeBp.route('/workflow/star', methods=['POST', 'OPTIONS'])
def star_workflow_api():
    """
    简历STAR工作流API

    请求参数:
        text: 简历文本内容

    返回:
        STAR分析结果
    """
    if request.method == 'OPTIONS':
        response = jsonify({'code': 200, 'message': 'OK'})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    try:
        data = request.get_json()

        # 获取请求参数
        text_content = data.get('text', '')

        # 导入STAR工作流模块
        from scripts.workflow.resume_star_workflow import analyze_resume_star

        result = analyze_resume_star(text_content)

        return jsonify({
            'code': 0,
            'message': '简历STAR分析成功',
            'data': result
        })
    
    except Exception as e:
        import traceback
        print(f"简历STAR分析错误: {str(e)}")
        print(traceback.format_exc())
        return jsonify({
            'code': 500,
            'message': f"简历STAR分析错误: {str(e)}",
            'data': None
        })